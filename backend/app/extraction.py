"""Document intake: raw text extraction, plus the precise "find a labeled
value in a span of text" primitive used by both extraction paths.

This module's `extract_fields` searches the whole document at once and is
kept as the dependency-free fallback. The primary path is
semantic_extraction.py, which narrows the search to one retrieved chunk of
the document before calling `match_value_in_text` here -- so a document's
title or table of contents mentioning a term's label can't be mistaken for
the line that actually states its value.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


@dataclass
class ExtractionCandidate:
    key_term_id: str
    label: str
    value: str
    confidence: float
    match_method: str = "regex"


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        return _extract_docx(content)
    # .txt and anything else: best-effort decode
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


_NUMBER_RE = re.compile(r"(-)?\$?\s*(-)?([\d,]+(?:\.\d+)?)")
_NUMBER_CLEAN_RE = re.compile(r"[^0-9.\-]")
_PERCENT_RE = re.compile(r"(-?[\d.]+)\s*%")
_DATE_RE = re.compile(
    r"([A-Za-z]{3,9}\.?\s+\d{1,2},?\s+\d{4}|\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{2,4})"
)


def _signed_digits(match: re.Match) -> str:
    """_NUMBER_RE matches a sign on either side of an optional '$' (a
    document might write a negative amount as "-$50,000" or "$-50,000") --
    this collapses whichever one fired into a single leading sign."""
    negative = bool(match.group(1)) or bool(match.group(2))
    digits = match.group(3)
    return f"-{digits}" if negative else digits


def _format_currency(digits: str) -> str:
    """Formats a signed digit string as canonical currency text -- '$',
    comma thousands separators, cents only when the source actually had a
    fractional part -- by parsing it as a number and reformatting it, not
    by trusting whatever grouping (or lack of one) the source document
    happened to use. "11600000", "11,600,000", and "11600000.00" all become
    the same "$11,600,000" (or "$11,600,000.00" if a fraction was present)."""
    negative = digits.startswith("-")
    raw = (digits[1:] if negative else digits).replace(",", "")
    value = float(raw)
    body = f"{value:,.2f}" if "." in raw else f"{value:,.0f}"
    return f"{'-' if negative else ''}${body}"


def typed_value(data_type: str, segment: str) -> tuple[str, bool]:
    """Returns (value, matched_with_expected_type)."""
    if data_type == "number":
        m = _NUMBER_RE.search(segment)
        if m:
            # Strip formatting (commas, etc.) so a "number" field's value is
            # ready to hand straight to a numeric mapping -- no ',' or other
            # non-numeric characters, unlike "currency" which keeps its
            # human-readable form (e.g. "11,600,000") for display.
            return _NUMBER_CLEAN_RE.sub("", _signed_digits(m)), True
    elif data_type == "currency":
        m = _NUMBER_RE.search(segment)
        if m:
            # Reformatted through _format_currency, not just $-prefixed --
            # a "currency" field's value should always read as "$11,600,000"
            # regardless of whether the source document wrote "11600000",
            # "11,600,000", or "$11,600,000.00".
            return _format_currency(_signed_digits(m)), True
    elif data_type == "percent":
        m = _PERCENT_RE.search(segment)
        if m:
            return f"{m.group(1)}%", True
    elif data_type == "date":
        m = _DATE_RE.search(segment)
        if m:
            return m.group(1), True
    fallback = segment.strip()
    return (fallback[:80], False) if fallback else ("", False)


def match_value_in_text(text: str, label: str, aliases: list[str], data_type: str) -> tuple[str, float, str, bool]:
    """Searches `text` for `label` (or one of `aliases`) and extracts a typed
    value from what follows it. Returns (value, confidence, match_method,
    typed_ok); value is "" and confidence 0.0 when nothing usable is found.

    `typed_ok` tells the caller whether the value actually matched the
    term's expected format (a number for "currency", a date pattern for
    "date", ...) versus being an untyped fallback grab of whatever text
    followed the label -- semantic_extraction.py uses this to keep looking
    at other chunks rather than settle for a currency field's value being
    some unrelated sentence fragment.

    This is the precise, single-span extraction primitive: it doesn't care
    whether `text` is a whole document or one retrieved chunk of one --
    semantic_extraction.py calls it against a single scoped chunk, while
    extract_fields below calls it against the full document text.
    """
    phrases = sorted([label, *aliases], key=len, reverse=True)
    phrases = [p for p in phrases if p.strip()]
    if not phrases:
        return "", 0.0, "regex", False

    alt = "|".join(re.escape(p) for p in phrases)
    pattern = re.compile(rf"(?:{alt})\s*[:\-–]?\s*(.{{0,100}})", re.IGNORECASE)
    match = pattern.search(text)
    if not match:
        return "", 0.0, "regex", False

    segment = match.group(1).split("\n")[0]
    exact_label_hit = match.group(0).lower().startswith(label.lower())

    value, typed_ok = typed_value(data_type, segment)
    if not value:
        return "", 0.0, "regex", False

    confidence = 0.9 if exact_label_hit else 0.7
    if not typed_ok:
        confidence -= 0.25
    confidence = max(0.0, min(confidence, 0.99))
    return value, round(confidence, 2), "regex", typed_ok


def extract_fields(text: str, key_terms: list[dict]) -> list[ExtractionCandidate]:
    """key_terms: list of {id, label, aliases, data_type}. Whole-document
    regex extraction -- no semantic chunking or retrieval. Kept as the
    dependency-free fallback semantic_extraction.py reaches for when a
    document has no useful embedding index (see extract_key_terms_semantic)."""
    results = []
    for term in key_terms:
        value, confidence, method, _typed_ok = match_value_in_text(
            text, term["label"], term.get("aliases", []), term.get("data_type", "text")
        )
        results.append(ExtractionCandidate(term["id"], term["label"], value, confidence, method))
    return results
